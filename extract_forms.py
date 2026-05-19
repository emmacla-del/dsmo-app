from docx import Document
from docx.table import Table
import os

docs_path = r"C:\Users\win\dsmo_app\docs\official_forms"
files = [
    "Questionnaire_ENTREPRISES ONEFOP .docx",
    "Questionnaire_Coopérative ONEFOP .docx",
    "Questionnaire_CTD ONEFOP .docx",
    "Questionnaire_ONG ONEFOP .docx"
]

for filename in files:
    filepath = os.path.join(docs_path, filename)
    print(f"\n{'='*80}")
    print(f"FILE: {filename}")
    print(f"{'='*80}")
    
    try:
        doc = Document(filepath)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        print(f"\nCONTENT PREVIEW (first 100 paragraphs):")
        print("-" * 80)
        
        for i, para in enumerate(doc.paragraphs[:100]):
            text = para.text.strip()
            if text:
                print(f"{i}: {text[:100]}")
        
        print(f"\n\nTABLE SUMMARY:")
        print("-" * 80)
        for t_idx, table in enumerate(doc.tables):
            print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
            if len(table.rows) > 0:
                print(f"  Row 0 (headers): {[cell.text.strip()[:30] for cell in table.rows[0].cells]}")
                
    except Exception as e:
        print(f"Error processing {filename}: {e}")
        import traceback
        traceback.print_exc()
