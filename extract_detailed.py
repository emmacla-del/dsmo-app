import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from pathlib import Path

docs_path = Path(r"C:\Users\win\dsmo_app\docs\official_forms")
docx_files = sorted(docs_path.glob("*.docx"))

for filepath in docx_files:
    fname = filepath.name
    print(f"\n{'='*100}")
    print(f"FILE: {fname}")
    print(f"{'='*100}\n")
    
    try:
        doc = Document(filepath)
        
        # Extract all text content
        print("ALL PARAGRAPHS WITH TEXT:")
        print("-" * 100)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"{i}: {text}")
        
        # Extract detailed table content
        print(f"\n\nTABLE DETAILS:")
        print("-" * 100)
        for t_idx, table in enumerate(doc.tables):
            print(f"\nTABLE {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols)")
            print("-" * 100)
            for r_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip()[:50] for cell in row.cells]
                print(f"R{r_idx}: {row_data}")
                
    except Exception as e:
        print(f"Error processing {filepath.name}: {e}")
        import traceback
        traceback.print_exc()

print("\n" + "="*100)
print("EXTRACTION COMPLETE")
print("="*100)
